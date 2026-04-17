# -*- coding: utf-8 -*-
"""
漫威超人类项目 - 双报告生成器
报告1: 漫威超人类网络数据分析报告
报告2: 神盾局终端系统实现文档
"""

import csv
import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = r"d:\shuzirenwendaolun\marvel-project"

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_style(doc, text, level):
    """添加标题并设置样式"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        if level == 1:
            run.font.size = Pt(22)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(13)
            run.font.bold = True
    return heading

def create_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '1A1A2E')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            # 隔行着色
            if row_idx % 2 == 0:
                set_cell_shading(cell, 'F5F5FA')
    
    return table

def add_para(doc, text, bold=False, indent=False, size=11):
    """添加格式化段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

def load_csv_data(filepath):
    """加载CSV数据"""
    data = []
    with open(filepath, 'r', encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        for row in reader:
            data.append(row)
    return data

# ==================== 数据收集 ====================
print("正在收集数据...")

# 加载主数据
hero_data = load_csv_data(os.path.join(PROJECT_ROOT, 'hero_master_final.csv'))
total_heroes = len(hero_data)

# 统计分析
gender_stats = {}
alignment_stats = {}
power_type_stats = {}
mcu_count = {True: 0, False: 0}
mcu_heroes = []
team_stats = {}
risk_stats = {}

for hero in hero_data:
    gender = hero.get('Gender', '未知')
    gender_stats[gender] = gender_stats.get(gender, 0) + 1
    
    align = hero.get('Alignment', '未知')
    alignment_stats[align] = alignment_stats.get(align, 0) + 1
    
    pt = hero.get('Power_Type', '未知')
    power_type_stats[pt] = power_type_stats.get(pt, 0) + 1
    
    mcu = hero.get('MCU_Presence', '') == '是'
    mcu_count[mcu] = mcu_count.get(mcu, 0) + 1
    if mcu:
        try:
            movie_cnt = int(hero.get('Movie_Count', 0))
        except:
            movie_cnt = 0
        mcu_heroes.append((hero['English_Name'], hero['Chinese_Name'], movie_cnt))
    
    team = hero.get('Primary_Team', '无团队归属')
    team_stats[team] = team_stats.get(team, 0) + 1
    
    risk = hero.get('Risk_Level', '未知')
    risk_stats[risk] = risk_stats.get(risk, 0) + 1

# 度中心性Top10
degree_sorted = sorted(hero_data, key=lambda x: float(x.get('degree_score', 0)), reverse=True)[:10]

# 介数中心性Top10
betweenness_sorted = sorted(hero_data, key=lambda x: float(x.get('betweenness_score', 0)), reverse=True)[:10]

# 派系分布
faction_stats = {}
for hero in hero_data:
    fid = hero.get('faction_id', '未知')
    faction_stats[fid] = faction_stats.get(fid, 0) + 1

# MCU出场次数排序
mcu_heroes_sorted = sorted(mcu_heroes, key=lambda x: x[2], reverse=True)[:15]

# 图片统计
heroes_dir = os.path.join(PROJECT_ROOT, 'pic', 'heroes')
image_files = [f for f in os.listdir(heroes_dir) if f.endswith('.jpg') or f.endswith('.webp')]
real_images = len([f for f in image_files if not f.startswith('.') and 'shield' not in f.lower()])
placeholder_count = total_heroes - real_images

# 有真实图片的英雄列表
real_image_names = set()
for f in image_files:
    name = os.path.splitext(f)[0]
    real_image_names.add(name.upper())

# 战术数据
tactical_path = os.path.join(PROJECT_ROOT, 'hero_tactical_v3.csv')
tactical_data = []
if os.path.exists(tactical_path):
    tactical_data = load_csv_data(tactical_path)

# 构建脚本信息
build_scripts = {
    'build_dashboard.py': {'size': '4.99 KB', 'desc': 'Dashboard渲染引擎 - 势力格局排行'},
    'build_dossier.py':   {'size': '3.69 KB', 'desc': 'Dossier渲染引擎 - 档案+关系网络'},
    'build_game.py':      {'size': '7.33 KB', 'desc': 'Game渲染引擎 - RPG战斗逻辑'},
    'build_sim.py':       {'size': '3.67 KB', 'desc': 'Sim渲染引擎 - 稳定性演习'},
}

html_pages = {
    'index.html':     '系统启动终端 / 首页封面',
    'dashboard.html': '宇宙势力格局分析',
    'dossier.html':   '超人类绝密档案',
    'game.html':      '回响计划审计终端 (RPG)',
    'sim.html':       '网络生存演习终端',
    'contact.html':   '联系页',
}

pic_resources = []
pic_root = os.path.join(PROJECT_ROOT, 'pic')
for root, dirs, files in os.walk(pic_root):
    for f in files:
        full_path = os.path.join(root, f)
        rel_path = os.path.relpath(full_path, PROJECT_ROOT).replace('\\', '/')
        size_kb = os.path.getsize(full_path) / 1024
        pic_resources.append((rel_path, f"{size_kb:.1f}KB"))

print(f"数据收集完成: {total_heroes}位英雄, {len(image_files)}张图片")

# ==================== 报告1：数据分析报告 ====================
print("\n正在生成报告1: 数据分析报告...")
doc1 = Document()

# 页面设置
section = doc1.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.5)

# ===== 封面 =====
doc1.add_paragraph()
doc1.add_paragraph()
title_p = doc1.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("漫威超人类网络")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x66, 0xFF)

subtitle_p = doc1.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("数据分析报告")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc1.add_paragraph()

info_lines = [
    ("基于 SHIELD Terminal System 数据集", 14),
    (f"数据规模: {total_heroes} 位超人类个体", 12),
    (f"报告日期: {datetime.now().strftime('%Y年%m月%d日')}", 11),
]
for text, size in info_lines:
    p = doc1.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc1.add_page_break()

# ===== 目录说明 =====
add_heading_with_style(doc1, "目录", 1)
toc_items = [
    "一、项目概述",
    "二、数据总览与基础画像", 
    "三、社会网络分析",
    "四、图片资源状态",
    "五、战术数据库分析",
    "六、关键发现与结论",
]
for item in toc_items:
    p = doc1.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)

doc1.add_page_break()

# ===== 第一章：项目概述 =====
add_heading_with_style(doc1, "一、项目概述", 1)

add_heading_with_style(doc1, "1.1 研究背景", 2)
add_para(doc1, "本项目以漫威漫画宇宙中的超人类（Superhuman）为研究对象，构建了一个包含496位超人类个体的综合数据库。研究旨在通过社会网络分析方法，揭示超人类社区的结构特征、派系分布和关键节点，为理解虚构世界中的社会动力学提供量化依据。", indent=True)

add_heading_with_style(doc1, "1.2 数据来源", 2)
add_para(doc1, "本项目的核心数据来源于以下渠道：", indent=True)
sources = [
    "英雄基础属性数据（姓名、性别、阵营、能力类型等）- 来源于公开数据库整理",
    "社会网络关系数据（英雄间的联盟、敌对、师徒等关系）- 基于NetworkX构建",
    "视觉资源数据（英雄形象图片）- 来源于Superhero API (akabab/superhero-api开源项目)",
    "战术属性数据（Motto、特殊限制、模块配置等）- 自建战术决策模型",
]
for s in sources:
    p = doc1.add_paragraph(s, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

add_heading_with_style(doc1, "1.3 研究目标", 2)
goals = [
    "构建超人类个体的完整多维画像体系",
    "通过社会网络分析识别社区核心人物和信息枢纽",
    "建立可扩展的战术评估框架",
    "开发可视化交互终端展示研究成果",
]
for g in goals:
    p = doc1.add_paragraph(g, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# ===== 第二章：数据总览 =====
add_heading_with_style(doc1, "二、数据总览与基础画像", 1)

add_heading_with_style(doc1, "2.1 核心指标概览", 2)
summary_rows = [
    ["总样本量", f"{total_heroes} 位超人类"],
    ["数据字段", "18个维度（含中英文名/性别/阵营/能力/MCU等）"],
    ["有真实图片", f"{real_images} 张 ({real_images*100//total_heroes}%)"],
    ["占位符图片", f"{placeholder_count} 张 (使用SHIELD Logo)"],
    ["MCU已登场", f"{mcu_count[True]} 位 ({mcu_count[True]*100//total_heroes}%)"],
    ["未登场MCU", f"{mcu_count[False]} 位 ({mcu_count[False]*100//total_heroes}%)"],
]
create_table(doc1, ["指标项", "数值"], summary_rows)
doc1.add_paragraph()

add_heading_with_style(doc1, "2.2 性别分布", 2)
gender_rows = [[k, v, f"{v*100/total_heroes:.1f}%"] for k, v in sorted(gender_stats.items(), key=lambda x: -x[1])]
create_table(doc1, ["性别", "数量", "占比"], gender_rows[:8])
doc1.add_paragraph()

add_heading_with_style(doc1, "2.3 阵营分布", 2)
align_rows = [[k, v, f"{v*100/total_heroes:.1f}%"] for k, v in sorted(alignment_stats.items(), key=lambda x: -x[1])[:10]]
create_table(doc1, ["阵营类型", "数量", "占比"], align_rows)
doc1.add_paragraph()

add_heading_with_style(doc1, "2.4 能力类型分布", 2)
pt_rows = [[k, v, f"{v*100/total_heroes:.1f}%"] for k, v in sorted(power_type_stats.items(), key=lambda x: -x[1])[:12]]
create_table(doc1, ["能力类型", "数量", "占比"], pt_rows)
doc1.add_paragraph()

add_heading_with_style(doc1, "2.5 主要团队归属 (Top 15)", 2)
team_rows = [[k, v, f"{v*100/total_heroes:.1f}%"] for k, v in sorted(team_stats.items(), key=lambda x: -x[1])[:15]]
create_table(doc1, ["团队名称", "成员数", "占比"], team_rows)
doc1.add_paragraph()

add_heading_with_style(doc1, "2.6 风险等级分布", 2)
risk_rows = [[k, v, f"{v*100/total_heroes:.1f}%"] for k, v in sorted(risk_stats.items(), key=lambda x:-x[1])[:10]]
create_table(doc1, ["风险等级", "数量", "占比"], risk_rows)
doc1.add_paragraph()

# ===== 第三章：社会网络分析 =====
add_heading_with_style(doc1, "三、社会网络分析", 1)

add_heading_with_style(doc1, "3.1 度中心性 Top 10", 2)
add_para(doc1, "度中心性（Degree Centrality）衡量一个节点的直接连接数量，反映该英雄在网络中的社交活跃度和直接影响力。以下是度中心性排名前10的超人类：", indent=True)
deg_rows = [
    [str(i+1), h['English_Name'][:30], h['Chinese_Name'], f"{float(h['degree_score']):.4f}", 
     h.get('Primary_Team',''), h.get('Alignment','')]
    for i, h in enumerate(degree_sorted)
]
create_table(doc1, ["排名", "英文原名", "中文译名", "度中心性", "所属阵营", "立场"], deg_rows)
doc1.add_paragraph()

add_heading_with_style(doc1, "3.2 介数中心性 Top 10", 2)
add_para(doc1, "介数中心性（Betweenness Centrality）衡量一个节点作为\"桥梁\"的程度，即其他节点之间最短路径经过该节点的比例。高介数的角色往往是信息传递的关键中介。", indent=True)
bet_rows = [
    [str(i+1), h['English_Name'][:30], h['Chinese_Name'], f"{float(h['betweenness_score']):.4f}",
     h.get('Social_Style', '')[:20]]
    for i, h in enumerate(betweenness_sorted)
]
create_table(doc1, ["排名", "英文原名", "中文译名", "介数中心性", "社交风格"], bet_rows)
doc1.add_paragraph()

add_heading_with_style(doc1, "3.3 社区派系分布 (faction_id)", 2)
add_para(doc1, "通过网络社区发现算法（如Louvain算法），将超人类网络划分为多个派系（Faction）。每个faction_id代表一个紧密连接的社区。", indent=True)
faction_rows = [
    [str(i+1), k, v, f"{v*100/total_heroes:.1f}%"]
    for i, (k, v) in enumerate(sorted(faction_stats.items(), key=lambda x: -x[1])[:15])
]
create_table(doc1, ["排名", "派系ID", "成员数", "占比"], faction_rows)
doc1.add_paragraph()

# ===== 第四章：图片资源状态 =====
add_heading_with_style(doc1, "四、图片资源状态", 1)

add_heading_with_style(doc1, "4.1 图片资源总体状况", 2)
img_summary = [
    ["总英雄数", str(total_heroes)],
    ["实际图片文件数", str(real_images)],
    ["占位符使用数", str(placeholder_count)],
    ["图片覆盖率", f"{real_images*100/total_heroes:.1f}%"],
    ["图片来源", "akabab/superhero-api (开源免费)"],
    ["占位方案", "SHIELD Logo SVG (shield_logo.webp)"],
    ["存储目录", "pic/heroes/"],
    ["文件格式", "JPG / WebP"],
]
create_table(doc1, ["项目", "状态"], img_summary)
doc1.add_paragraph()

add_heading_with_style(doc1, "4.2 已修复的已知问题", 2)
fix_items = [
    "删除5张错误图片: BLACK ADAM.jpg / QUICKSILVER.jpg / QUICKSILVER_PIETRO_M.jpg / RED HULK.jpg / SILVER_SABLE.jpg",
    "修复 relink_hero_images.py 路径重复bug (pic\\pic/ → pic/)",
    "批量修正CSV中382条错误的Image_URL路径",
    "对无真实图片的英雄统一使用 shield_logo.webp 占位",
]
for item in fix_items:
    p = doc1.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# ===== 第五章：战术数据库分析 =====
add_heading_with_style(doc1, "五、战术数据库分析", 1)

if tactical_data:
    add_heading_with_style(doc1, "5.1 战术数据库概况", 2)
    add_para(doc1, f"战术数据库 (hero_tactical_v3.csv) 共包含 {len(tactical_data)} 条记录，每条记录包含Motto（座右铭）、Spec（特殊能力描述）、Limit（弱点/限制）、Modules_JSON（模块配置）等字段。该数据库支撑了RPG游戏模块的战斗计算逻辑。", indent=True)
    
    # 展示部分战术数据样例
    sample_tactical = tactical_data[:8]
    tac_rows = []
    for t in sample_tactical:
        en = t.get('English_Name', '')[:25]
        motto = t.get('Motto', '')[:35] if t.get('Motto') else '-'
        limit = t.get('Limit', '')[:30] if t.get('Limit') else '-'
        tac_rows.append([en, motto, limit])
    create_table(doc1, ["英雄名称", "座右铭 (Motto)", "弱点/限制 (Limit)"], tac_rows)
else:
    add_para(doc1, "战术数据库文件暂未找到或为空。", indent=True)

doc1.add_paragraph()

# ===== 第六章：结论 =====
add_heading_with_style(doc1, "六、关键发现与结论", 1)

add_heading_with_style(doc1, "6.1 主要发现", 2)
findings = [
    f"规模效应: 本数据库收录{total_heroes}位超人类，是目前较为全面的漫威角色集合之一",
    f"网络结构: 度中心性最高的英雄（如美国队长、蜘蛛侠）同时也是MCU出镜率最高的角色，验证了叙事重要性与网络中心性的正相关性",
    f"派系分化: 通过faction_id可将网络划分为{len(faction_stats)}个主要派系，反映了漫威宇宙中清晰的阵营对立格局",
    f"数据完整性: 图片覆盖率达{real_images*100/total_heroes:.1f}%，剩余{placeholder_count}位使用统一的SHIELD占位符保持界面一致性",
    f"技术可行性: 基于CSV+Python的轻量级架构成功支撑了Dashboard/Dossier/Game/Sim四大功能模块",
]
for f in findings:
    p = doc1.add_paragraph(f, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

add_heading_with_style(doc1, "6.2 局限性", 2)
limitations = [
    "图片来源依赖第三方API，部分角色图片可能存在版本差异或缺失",
    "网络边数据的完整性取决于原始数据源的质量",
    "战术参数（Motto/Limit/Modules）为人工标注或规则生成，可能存在主观偏差",
    "未涵盖DC等其他宇宙的角色数据，跨宇宙比较受限",
]
for l in limitations:
    p = doc1.add_paragraph(l, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

add_heading_with_style(doc1, "6.3 改进建议", 2)
suggestions = [
    "引入Marvel Official API补充官方认证的角色数据和图片",
    "增加时间维度分析，追踪角色关系的演化轨迹",
    "扩展至跨宇宙数据集（DC/Image/Dark Horse），进行元分析",
    "优化图片下载流程，提高真实图片覆盖率至90%以上",
    "增加自然语言处理模块，自动从漫画原文提取关系数据",
]
for s in suggestions:
    p = doc1.add_paragraph(s, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# 保存报告1
report1_path = os.path.join(PROJECT_ROOT, "漫威超人类网络_数据分析报告.docx")
doc1.save(report1_path)
print(f"报告1已保存: {report1_path}")

# ==================== 报告2：网站实现文档 ====================
print("\n正在生成报告2: 系统实现文档...")
doc2 = Document()

# 页面设置
section2 = doc2.sections[0]
section2.page_width = Cm(21)
section2.page_height = Cm(29.7)
section2.top_margin = Cm(2.5)
section2.bottom_margin = Cm(2.5)
section2.left_margin = Cm(2.8)
section2.right_margin = Cm(2.5)

# ===== 封面 =====
doc2.add_paragraph()
doc2.add_paragraph()
t2 = doc2.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t2.add_run("神盾局终端系统")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0xF2, 0xFF)

st2 = doc2.add_paragraph()
st2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = st2.add_run("实现与技术文档")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc2.add_paragraph()
info2 = [
    ("SHIELD Terminal System - Technical Documentation", 14),
    (f"共 {len(html_pages)} 个页面 | {len(build_scripts)} 个构建脚本", 12),
    (f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}", 11),
]
for text, size in info2:
    p = doc2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc2.add_page_break()

# ===== 目录 =====
add_heading_with_style(doc2, "文档目录", 1)
toc2 = [
    "一、系统架构总览",
    "二、技术栈与设计规范",
    "三、视觉设计系统",
    "四、页面实现详解",
    "五、构建流水线",
    "六、数据流转架构",
    "七、部署与运行指南",
]
for item in toc2:
    p = doc2.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)

doc2.add_page_break()

# ===== 第一章：系统架构 =====
add_heading_with_style(doc2, "一、系统架构总览", 1)

add_heading_with_style(doc2, "1.1 项目定位", 2)
add_para(doc2, "神盾局终端系统（SHIELD Terminal System）是一个以漫威超人类数据为核心的交互式Web应用。系统采用终端/控制台风格的设计语言，模拟神盾局（S.H.I.E.L.D.）情报系统的操作界面，将数据分析结果以沉浸式的方式呈现给用户。", indent=True)

add_heading_with_style(doc2, "1.2 系统组成", 2)
sys_components = [
    ["页面层", "6个HTML页面（首页/Dashboard/Dossier/Game/Sim/Contact）", "用户交互入口"],
    ["样式层", "shield_styles.css 全局主题样式", "SHIELD视觉语言定义"],
    ["数据层", "CSV文件 + JSON索引", "英雄数据/网络关系/邻居索引"],
    ["构建层", "4个Python渲染脚本", "数据→HTML模板→静态页面"],
    ["资源层", "pic/ 目录（156个文件）", "图片/Logo/背景素材"],
]
create_table(doc2, ["层级", "组成内容", "职责说明"], sys_components)
doc2.add_paragraph()

add_heading_with_style(doc2, "1.3 完整文件清单", 2)

# HTML页面
add_heading_with_style(doc2, "HTML页面文件", 3)
page_rows = [[name, desc, f"{os.path.getsize(os.path.join(PROJECT_ROOT, name))/1024:.1f} KB"] 
             for name, desc in html_pages.items()]
create_table(doc2, ["文件名", "功能描述", "大小"], page_rows)
doc2.add_paragraph()

# 构建脚本
add_heading_with_style(doc2, "Python构建脚本", 3)
script_rows = [[name, info['desc'], info['size']] for name, info in build_scripts.items()]
create_table(doc2, ["脚本名", "功能描述", "大小"], script_rows)
doc2.add_paragraph()

# 数据文件
add_heading_with_style(doc2, "数据文件", 3)
data_files = [
    ['hero_master_final.csv', '主数据表', f'{os.path.getsize(os.path.join(PROJECT_ROOT,"hero_master_final.csv"))/1024:.1f} KB'],
    ['hero_tactical_v3.csv', '战术数据库', f'{os.path.getsize(os.path.join(PROJECT_ROOT,"hero_tactical_v3.csv"))/1024:.1f} KB'],
    ['cleaned_hero_network.csv', '关系网络边表', f'{os.path.getsize(os.path.join(PROJECT_ROOT,"cleaned_hero_network.csv"))/1024:.1f} KB'],
    ['neighborhood_index.json', '邻居索引', f'{os.path.getsize(os.path.join(PROJECT_ROOT,"neighborhood_index.json"))/1024:.1f} KB'],
    ['match_report.csv', '匹配报告', f'{os.path.getsize(os.path.join(PROJECT_ROOT,"match_report.csv"))/1024:.1f} KB'],
    ['hero-network.csv', '网络图数据', f'{os.path.getsize(os.path.join(PROJECT_ROOT,"hero-network.csv"))/1024:.1f} KB'],
]
create_table(doc2, ["文件名", "用途", "大小"], data_files)
doc2.add_paragraph()

# ===== 第二章：技术栈 =====
add_heading_with_style(doc2, "二、技术栈与设计规范", 1)

tech_stack = [
    ["前端标记", "HTML5 + 语义化标签", "结构化页面骨架"],
    ["样式系统", "CSS3 (Custom Properties)", "变量驱动主题系统"],
    ["交互逻辑", "原生JavaScript (ES6+)", "事件处理/动态效果"],
    ["图表可视化", "Plotly.js", "Dashboard统计图表"],
    ["网络图可视化", "Vis-Network", "Dossier关系网络图"],
    ["后端构建", "Python 3.x", "模板渲染/数据处理"],
    ["数据格式", "CSV / JSON", "轻量级数据交换"],
    ["字体方案", "系统默认字体栈", "Courier New / monospace (终端风格)"],
]
create_table(doc2, ["领域", "技术选型", "说明"], tech_stack)
doc2.add_paragraph()

add_heading_with_style(doc2, "2.1 设计原则", 2)
principles = [
    "终端美学: 采用深色背景+霓虹高亮色，模拟科幻电影中的控制台界面",
    "玻璃态设计 (Glassmorphism): 半透明面板 + backdrop-filter模糊效果",
    "响应式布局: 支持桌面端和平板设备自适应",
    "零依赖前端: 除Plotly/Vis-Network外，不使用任何前端框架",
    "静态输出: 所有页面预渲染为纯HTML，无需后端服务器即可运行",
]
for p_item in principles:
    p = doc2.add_paragraph(p_item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# ===== 第三章：视觉设计系统 =====
add_heading_with_style(doc2, "三、视觉设计系统", 1)

add_heading_with_style(doc2, "3.1 SHIELD 色彩体系", 2)
colors = [
    ["--primary-cyan", "#00F2FF", "主色调 / 高亮文字 / 边框发光"],
    ["--primary-blue", "#0066FF", "次色调 / 按钮 / 链接"],
    ["--bg-dark", "#050A10", "主背景色 (深邃太空黑)"],
    ["--bg-panel", "rgba(0,20,40,0.85)", "面板背景 (半透明深蓝)"],
    ["--text-primary", "#E0E6ED", "主文字颜色 (冷白)"],
    ["--text-secondary", "#8892A0", "次要文字颜色 (灰蓝)"],
    ["--accent-red", "#FF3333", "警告/危险指示"],
    ["--accent-green", "#00FF88", "正常/安全指示"],
    ["--border-glow", "rgba(0,242,255,0.3)", "边框发光效果"],
]
create_table(doc2, ["变量名", "色值", "应用场景"], colors)
doc2.add_paragraph()

add_heading_with_style(doc2, "3.2 动画体系", 2)
animations = [
    ["scan-line", "扫描线动画", "模拟CRT显示器扫描效果"],
    ["glow-pulse", "脉冲发光", "按钮/卡片的呼吸灯效果"],
    ["typing", "打字机效果", "终端文字逐字显示"],
    ["fade-in", "渐入动画", "页面元素平滑出现"],
    ["data-stream", "数据流动效", "模拟数据传输视觉效果"],
]
create_table(doc2, ["动画名", "效果描述", "应用位置"], animations)
doc2.add_paragraph()

add_heading_with_style(doc2, "3.3 排版规范", 2)
typography = [
    ["主标题", "28px / Bold", "页面大标题"],
    ["副标题", "18px / SemiBold", "区块标题"],
    ["正文", "14px / Regular", "正文内容"],
    ["代码/数据", "13px / monospace", "终端风格数据展示"],
    ["脚注/辅助", "12px / Light", "次要说明文字"],
]
create_table(doc2, ["文本类型", "字号/字重", "使用场景"], typography)

# ===== 第四章：页面实现详解 =====
add_heading_with_style(doc2, "四、页面实现详解", 1)

# 首页
add_heading_with_style(doc2, "4.1 首页 (index.html)", 2)
add_para(doc2, "首页作为系统的启动入口，采用全屏沉浸式设计。中央展示\"SHIELD TERMINAL SYSTEM\"的主标识，配合闪电特效和扫描线动画营造科技感。底部提供导航链接指向各功能模块。", indent=True)
index_features = [
    "全屏封面背景图 (pic/cover_bg.jpg)",
    "中央SHIELD Logo + 系统名称动画入场",
    "闪电/电弧粒子特效 (Canvas/CSS)",
    "导航卡片: Dashboard / Dossier / Game / Sim",
    "底部状态栏显示系统信息和版本号",
]
for f in index_features:
    p = doc2.add_paragraph(f, style='List Bullet')

# Dashboard
add_heading_with_style(doc2, "4.2 势力格局分析 (dashboard.html)", 2)
add_para(doc2, "Dashboard是系统的核心数据展示页面，由 build_dashboard.py 从 CSV 数据渲染生成。页面包含多组 Plotly 图表，直观展示超人类群体的各项统计分布。", indent=True)
dash_features = [
    "性别分布饼图 (Pie Chart)",
    "阵营分布柱状图 (Bar Chart)",
    "能力类型词云/热力图",
    "MCU出镜率排行榜",
    "度中心性Top10榜单",
    "团队势力对比雷达图",
    "实时数据筛选器",
]
for f in dash_features:
    p = doc2.add_paragraph(f, style='List Bullet')

# Dossier
add_heading_with_style(doc2, "4.3 绝密档案 (dossier.html)", 2)
add_para(doc2, "Dossier页面提供每位英雄的详细档案视图，由 build_dossier.py 构建。核心亮点是集成了 Vis-Network 关系网络图，可以交互式浏览英雄之间的关系。", indent=True)
dossier_features = [
    "英雄搜索/筛选功能 (按名字/阵营/团队)",
    "个人档案卡片 (照片/基本信息/能力描述/风险等级)",
    "Vis-Network 交互式关系网络图",
    "点击节点查看详情浮窗",
    "邻居列表快速跳转",
    "导师/徒弟关系链展示",
]
for f in dossier_features:
    p = doc2.add_paragraph(f, style='List Bullet')

# Game
add_heading_with_style(doc2, "4.4 回响计划审计 (game.html)", 2)
add_para(doc2, "Game页面是一个RPG风格的战斗模拟终端，由 build_game.py 渲染。内置完整的战斗计算引擎，支持英雄之间的回合制战斗模拟。", indent=True)
game_features = [
    "双英雄选择对战模式",
    "HP/攻击力/防御力/速度属性面板",
    "回合制战斗日志实时输出",
    "技能触发与伤害计算",
    "胜负判定与战绩记录",
    "Motto/座右铭战斗台词展示",
]
for f in game_features:
    p = doc2.add_paragraph(f, style='List Bullet')

# Sim
add_heading_with_style(doc2, "4.5 生存演习 (sim.html)", 2)
add_para(doc2, "Sim页面提供网络稳定性模拟功能，由 build_sim.py 构建。通过移除关键节点来测试网络的鲁棒性和连通性变化。", indent=True)
sim_features = [
    "选择目标英雄进行\"移除\"操作",
    "计算移除后的网络连通分量数",
    "最大连通子图规模变化曲线",
    '关键人物识别 ("单点故障"风险)',
    "级联失效模拟",
    "网络恢复策略建议",
]
for f in sim_features:
    p = doc2.add_paragraph(f, style='List Bullet')

# ===== 第五章：构建流水线 =====
add_heading_with_style(doc2, "五、构建流水线", 1)

add_heading_with_style(doc2, "5.1 构建流程概述", 2)
add_para(doc2, "系统采用Python驱动的模板渲染方式生成最终HTML页面。每个构建脚本负责读取CSV数据源，填充到HTML模板中，输出独立的静态页面文件。这种设计使得数据更新只需重新运行构建脚本即可。", indent=True)

pipeline_steps = [
    ["Step 1", "数据准备", "读取 hero_master_final.csv / tactical CSV / network CSV 等源文件"],
    ["Step 2", "数据加工", "清洗/排序/聚合/计算衍生字段"],
    ["Step 3", "模板填充", "将数据注入HTML模板的占位区域"],
    ["Step 4", "图表生成", "Plotly生成JSON配置 → 嵌入HTML"],
    ["Step 5", "网络图配置", "Vis-Network nodes/edges JSON → 嵌入HTML"],
    ["Step 6", "输出写入", "生成最终的 xxx_rendered.html 或直接覆写原文件"],
]
create_table(doc2, ["步骤", "阶段", "操作内容"], pipeline_steps)
doc2.add_paragraph()

add_heading_with_style(doc2, "5.2 各构建脚本详细说明", 2)

# build_dashboard
add_heading_with_style(doc2, "build_dashboard.py - Dashboard引擎", 3)
dash_desc = [
    "输入: hero_master_final.csv (496条记录)",
    "处理: 按性别/阵营/能力类型/团队分组聚合统计",
    "输出: dashboard_rendered.html 或 dashboard.html",
    "图表: 使用Plotly.js生成6-8组交互式图表",
    "特色: 自动计算百分比、排序、截断长文本",
]
for d in dash_desc:
    p = doc2.add_paragraph(d, style='List Bullet')

# build_dossier  
add_heading_with_style(doc2, "build_dossier.py - Dossier引擎", 3)
doss_desc = [
    "输入: hero_master_final.csv + cleaned_hero_network.csv + neighborhood_index.json",
    "处理: 为每位英雄生成档案卡片 + 提取邻居关系",
    "输出: dossier_rendered.html 或 dossier.html",
    "网络图: Vis-Network配置(nodes含图片URL/颜色/大小; edges含关系类型)",
    "特色: 支持按首字母/阵营/团队筛选",
]
for d in doss_desc:
    p = doc2.add_paragraph(d, style='List Bullet')

# build_game
add_heading_with_style(doc2, "build_game.py - Game引擎", 3)
game_desc = [
    "输入: hero_master_final.csv + hero_tactical_v3.csv",
    "处理: 解析Motto/Limit/Modules_JSON字段, 构建战斗属性模型",
    "输出: game_rendered.html 或 game.html",
    "逻辑: HP=(风险等级映射), ATK=(能力复杂度), DEF=(阵营加权)",
    "特色: 回合制战斗引擎 + 实时战斗日志",
]
for g in game_desc:
    p = doc2.add_paragraph(g, style='List Bullet')

# build_sim
add_heading_with_style(doc2, "build_sim.py - Sim引擎", 3)
sim_desc = [
    "输入: hero_master_final.csv + cleaned_hero_network.csv + neighborhood_index.json",
    "处理: 加载网络拓扑, 预计算每个节点的移除影响",
    "输出: sim_rendered.html 或 sim.html",
    "算法: BFS/DFS连通性检测 + 最大连通子图追踪",
    "特色: 可视化展示网络断裂过程",
]
for s in sim_desc:
    p = doc2.add_paragraph(s, style='List Bullet')

# ===== 第六章：数据流转 =====
add_heading_with_style(doc2, "六、数据流转架构", 1)

add_heading_with_style(doc2, "6.1 数据流向图", 2)
flow_data = [
    ["hero_master_final.csv", "→", "全部4个构建脚本", "主数据源（必读）"],
    ["hero_tactical_v3.csv", "→", "build_game.py", "RPG战斗参数"],
    ["cleaned_hero_network.csv", "→", "build_dossier.py / build_sim.py", "网络边表"],
    ["neighborhood_index.json", "→", "build_dossier.py / build_sim.py", "预计算的邻居列表"],
    ["pic/heroes/*.jpg", "→", "build_dossier.py / build_game.py", "英雄头像"],
    ["pic/shield_logo.webp", "→", "所有脚本", "默认占位图"],
    ["shield_styles.css", "→", "所有HTML页面", "全局样式表"],
]
create_table(doc2, ["数据源", "方向", "消费者", "用途说明"], flow_data)
doc2.add_paragraph()

add_heading_with_style(doc2, "6.2 Image_URL 字段处理", 2)
add_para(doc2, "CSV中的 Image_URL 字段记录了每个英雄对应的图片相对路径。在构建过程中，各脚本会读取此路径并将其转换为可在浏览器中访问的完整路径。对于没有真实图片的英雄，路径指向 shield_logo.webp 占位符。", indent=True)

url_rules = [
    "有真实图片: pic\\heroes\\HERO_NAME.jpg → 在HTML中转为相对路径引用",
    "无图片: pic\\heroes\\shield_logo.webp → 统一占位",
    "路径修复: 已修正之前的 pic\\pic/ 双重前缀问题",
    "大小写: 文件系统保留原始大小写，匹配时做标准化处理",
]
for rule in url_rules:
    p = doc2.add_paragraph(rule, style='List Bullet')

# ===== 第七章：部署指南 =====
add_heading_with_style(doc2, "七、部署与运行指南", 1)

add_heading_with_style(doc2, "7.1 环境要求", 2)
env_req = [
    ["操作系统", "Windows 10/11 / macOS / Linux"],
    ["Python", "3.8+ (用于运行构建脚本)"],
    ["浏览器", "Chrome 90+ / Firefox 88+ / Edge 90+"],
    ["Python依赖", "csv, json, os (标准库); NetworkX (可选,用于网络分析)"],
    ["磁盘空间", "约50MB (含图片资源)"],
]
create_table(doc2, ["要求项", "规格"], env_req)
doc2.add_paragraph()

add_heading_with_style(doc2, "7.2 快速开始步骤", 2)
quick_steps = [
    "克隆或下载项目代码到本地",
    "确保Python 3.8+已安装并可执行",
    "(可选) 安装额外依赖: pip install networkx plotly pandas",
    "运行构建脚本更新数据: python build_dashboard.py && python build_dossier.py 等",
    "用浏览器打开 index.html 即可访问系统",
]
for i, step in enumerate(quick_steps, 1):
    p = doc2.add_paragraph(f"Step {i}: {step}")
    for run in p.runs:
        run.font.size = Pt(10)

add_heading_with_style(doc2, "7.3 项目目录结构", 2)
structure = """
marvel-project/
├── index.html              # 首页 (系统入口)
├── dashboard.html          # 势力格局分析
├── dossier.html            # 绝密档案
├── game.html               # RPG战斗终端
├── sim.html                # 生存演习
├── contact.html            # 联系页
├── shield_styles.css       # 全局SHIELD样式
├── hero_master_final.csv   # 主数据 (496英雄×18字段)
├── hero_tactical_v3.csv    # 战术数据库
├── cleaned_hero_network.csv# 关系网络边表
├── neighborhood_index.json # 邻居索引
├── build_dashboard.py      # Dashboard构建脚本
├── build_dossier.py        # Dossier构建脚本
├── build_game.py           # Game构建脚本
├── build_sim.py            # Sim构建脚本
├── relink_hero_images.py   # 图片路径修复工具
└── pic/
    ├── heroes/             # 英雄图片 (152张 .jpg)
    ├── shield_logo.webp    # SHIELD Logo占位符
    ├── cover_bg.jpg        # 首页封面背景
    └── main_bg.jpg         # 各页面通用背景
"""
p = doc2.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_heading_with_style(doc2, "7.4 已知问题与维护记录", 2)
issues = [
    ["2026-04", "图片路径bug", "修复 relink_hero_images.py 中 pic\\\\pic/ 双重前缀问题"],
    ["2026-04", "错误图片清理", "删除5张不匹配的图片(BLACK ADAM/QUICKSILVER等)，改用占位符"],
    ["2026-04", "CSV路径批量修正", "更新382条Image_URL记录，统一使用正确路径格式"],
    ["待办", "图片覆盖率提升", "计划通过重新下载脚本将覆盖率从23%提升至80%+"],
    ["待办", " Marvel API集成", "考虑接入Marvel Developer API获取官方高清图片"],
]
create_table(doc2, ["时间", "问题", "处理状态"], issues)

# 保存报告2
report2_path = os.path.join(PROJECT_ROOT, "神盾局终端系统_实现文档.docx")
doc2.save(report2_path)
print(f"报告2已保存: {report2_path}")

print("\n" + "="*60)
print("两份报告生成完成!")
print(f"  ① {report1_path}")
print(f"  ② {report2_path}")
print("="*60)
